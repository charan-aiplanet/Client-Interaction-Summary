import streamlit as st
import pandas as pd
import json
import time
import io
from datetime import datetime
from typing import Dict, List, Any
import asyncio
import threading
from dataclasses import dataclass
import PyPDF2
import docx
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from openai import AzureOpenAI
import autogen
from autogen import ConversableAgent, UserProxyAgent
import hashlib
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import uuid

# Page configuration
st.set_page_config(
    page_title="Client Interaction Summary Generator",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for cozy and dynamic UI
st.markdown("""
<style>
    .main {
        padding: 1rem;
    }
    
    .login-container {
        max-width: 400px;
        margin: 0 auto;
        padding: 2rem;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        border-radius: 20px;
        box-shadow: 0 10px 30px rgba(0,0,0,0.2);
        color: white;
        text-align: center;
        margin-top: 5rem;
    }
    
    .login-header {
        font-size: 2rem;
        margin-bottom: 2rem;
        font-weight: bold;
    }
    
    .agent-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        border-radius: 15px;
        padding: 1rem;
        margin: 0.5rem 0;
        color: white;
        box-shadow: 0 4px 15px rgba(0,0,0,0.1);
        transition: transform 0.3s ease;
    }
    
    .agent-card:hover {
        transform: translateY(-2px);
    }
    
    .processing-step {
        background: linear-gradient(45deg, #ff9a9e 0%, #fad0c4 100%);
        border-radius: 10px;
        padding: 1rem;
        margin: 0.5rem 0;
        border-left: 4px solid #ff6b6b;
        animation: fadeIn 0.5s ease-in;
    }
    
    .agent-interaction {
        background: linear-gradient(45deg, #a8edea 0%, #fed6e3 100%);
        border-radius: 10px;
        padding: 1rem;
        margin: 0.5rem 0;
        border: 2px solid #4ecdc4;
        animation: slideIn 0.5s ease-out;
    }
    
    .summary-output {
        background: linear-gradient(135deg, #ffecd2 0%, #fcb69f 100%);
        border-radius: 15px;
        padding: 1.5rem;
        margin: 1rem 0;
        box-shadow: 0 6px 20px rgba(0,0,0,0.1);
    }
    
    .history-item {
        background: linear-gradient(45deg, #e3f2fd 0%, #f3e5f5 100%);
        border-radius: 8px;
        padding: 0.8rem;
        margin: 0.3rem 0;
        border-left: 3px solid #2196f3;
        cursor: pointer;
        transition: all 0.3s ease;
    }
    
    .history-item:hover {
        transform: translateX(5px);
        box-shadow: 0 2px 10px rgba(0,0,0,0.1);
    }
    
    @keyframes fadeIn {
        from { opacity: 0; }
        to { opacity: 1; }
    }
    
    @keyframes slideIn {
        from { transform: translateX(-20px); opacity: 0; }
        to { transform: translateX(0); opacity: 1; }
    }
    
    .status-indicator {
        display: inline-block;
        width: 12px;
        height: 12px;
        border-radius: 50%;
        margin-right: 8px;
        animation: pulse 2s infinite;
    }
    
    .status-active { background-color: #4CAF50; }
    .status-waiting { background-color: #FF9800; }
    .status-complete { background-color: #2196F3; }
    
    @keyframes pulse {
        0% { transform: scale(1); }
        50% { transform: scale(1.2); }
        100% { transform: scale(1); }
    }
    
    .logout-btn {
        position: fixed;
        top: 10px;
        right: 10px;
        z-index: 1000;
    }
    .username-green {
        color: #4CAF50;
        font-weight: bold;
    }
</style>
""", unsafe_allow_html=True)

@dataclass
class AgentStatus:
    name: str
    status: str  # active, waiting, complete
    current_task: str
    messages: List[str]

@dataclass
class ProcessingHistory:
    id: str
    timestamp: datetime
    filename: str
    client_name: str
    summary: str
    full_results: Dict[str, Any]

class DocumentProcessor:
    """Handles document processing for different file formats"""
    
    @staticmethod
    def extract_text_from_pdf(file) -> str:
        try:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            st.error(f"Error processing PDF: {str(e)}")
            return ""
    
    @staticmethod
    def extract_text_from_docx(file) -> str:
        try:
            doc = docx.Document(file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            st.error(f"Error processing DOCX: {str(e)}")
            return ""
    
    @staticmethod
    def extract_text_from_txt(file) -> str:
        try:
            return file.read().decode('utf-8')
        except Exception as e:
            st.error(f"Error processing TXT: {str(e)}")
            return ""

class AzureOpenAIWrapper:
    """Wrapper for Azure OpenAI to work with AutoGen"""
    
    def __init__(self, api_key: str, endpoint: str, api_version: str = "2024-02-15-preview", model: str = "gpt-4"):
        self.client = AzureOpenAI(
            api_key=api_key,
            azure_endpoint=endpoint,
            api_version=api_version
        )
        self.model = model
    
    def generate_response(self, messages: List[Dict]) -> str:
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=messages,
                temperature=0.7,
                max_tokens=2048
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"Error generating response: {str(e)}"

class ClientSummaryAgents:
    """Multi-agent system for client interaction summary generation"""
    
    def __init__(self, azure_client: AzureOpenAIWrapper):
        self.azure_client = azure_client
        self.agent_statuses = {}
        self.interaction_log = []
        self.setup_agents()
    
    def setup_agents(self):
        """Initialize the multi-agent system"""
        
        # Document Analyzer Agent
        self.doc_analyzer = ConversableAgent(
            name="DocumentAnalyzer",
            system_message="""You are a Document Analyzer agent. Your role is to:
            1. Analyze the structure and content of client interaction documents
            2. Identify key sections, participants, and topics discussed
            3. Extract relevant information for summary generation
            4. Pass findings to the Summary Generator agent
            
            Be thorough and systematic in your analysis.""",
            llm_config={"config_list": [{"model": "azure", "api_key": "placeholder"}]},
            human_input_mode="NEVER"
        )
        
        # Summary Generator Agent
        self.summary_generator = ConversableAgent(
            name="SummaryGenerator",
            system_message="""You are an AI assistant designed to create concise and informative summaries of client interactions for Relationship Managers (RMs). Your task is to analyze the content of a client meeting or call and generate a structured summary.

You MUST use the following exact format for all client interaction summaries (WITHOUT any ** markdown formatting):

Client Interaction Summary
Date of Meeting: [Insert Date] 
Participants: [List Participants] 
Client Name: [Insert Client Name] 
Meeting Type: [Call/Meeting/Other]

1. Objectives of the Meeting : 
[Clearly state the purpose and goals of the meeting]

2. Key Discussion Points : 
[List the main topics, issues, and subjects discussed during the interaction]

3. Decisions Made : 
[Document any decisions, agreements, or resolutions reached during the meeting]

4. Action Items :
[List specific action items with responsible for RM and Client separately and timelines where mentioned]

Key Takeaways :
[Summarize the most important topics in bullet points, outcomes, or next steps from the meeting]

Instructions:
1. Always use the exact format structure shown above with the same headings but NO ** markdown formatting
2. Extract relevant information from the document and place it under the appropriate sections
3. If information for any section is not available in the document, note "[Information not available in document]"
4. For each section, provide clear, concise information about what was discussed
5. Highlight any important decisions, action items, or follow-up tasks
6. Keep the language professional and objective throughout the summary
7. Ensure the summary is concise while capturing all essential information
8. If any part of the interaction is unclear or seems to be missing context, note this in the relevant section
9. DO NOT use any ** formatting or markdown - output clean, plain text with proper headings

Take analyzed document content from the Document Analyzer and create structured, professional summaries following this exact format.""",
            llm_config={"config_list": [{"model": "azure", "api_key": "placeholder"}]},
            human_input_mode="NEVER"
        )
        
        # Quality Reviewer Agent
        self.quality_reviewer = ConversableAgent(
            name="QualityReviewer",
            system_message="""You are a Quality Reviewer agent. Your role is to:
            1. Review summaries generated by the Summary Generator
            2. Check for completeness, accuracy, and professional tone
            3. Ensure all key points and action items are captured
            4. Suggest improvements or approve the final summary
            5. Provide the final polished summary
            
            Be critical but constructive in your review.""",
            llm_config={"config_list": [{"model": "azure", "api_key": "placeholder"}]},
            human_input_mode="NEVER"
        )
        
        # Coordinator Agent
        self.coordinator = UserProxyAgent(
            name="Coordinator",
            system_message="You coordinate the summary generation process between agents.",
            human_input_mode="NEVER",
            code_execution_config=False
        )
        
        # Initialize agent statuses
        for agent_name in ["DocumentAnalyzer", "SummaryGenerator", "QualityReviewer", "Coordinator"]:
            self.agent_statuses[agent_name] = AgentStatus(
                name=agent_name,
                status="waiting",
                current_task="Initialized",
                messages=[]
            )
    
    def update_agent_status(self, agent_name: str, status: str, task: str, message: str = ""):
        """Update agent status and log interactions"""
        if agent_name in self.agent_statuses:
            self.agent_statuses[agent_name].status = status
            self.agent_statuses[agent_name].current_task = task
            if message:
                self.agent_statuses[agent_name].messages.append(message)
                self.interaction_log.append({
                    "timestamp": datetime.now().strftime("%H:%M:%S"),
                    "agent": agent_name,
                    "message": message,
                    "task": task
                })
    
    def process_document(self, document_text: str, format_template: str = "") -> Dict[str, Any]:
        """Process document through the multi-agent system"""
        
        # Step 1: Document Analysis
        self.update_agent_status("DocumentAnalyzer", "active", "Analyzing document structure")
        
        analysis_prompt = f"""
        Analyze this client interaction document and identify:
        1. Participants involved
        2. Main topics discussed
        3. Key decisions made
        4. Action items mentioned
        5. Important dates or deadlines
        6. Overall meeting context
        
        Document content:
        {document_text}
        """
        
        analysis_messages = [{"role": "user", "content": analysis_prompt}]
        analysis_result = self.azure_client.generate_response(analysis_messages)
        
        self.update_agent_status("DocumentAnalyzer", "complete", "Document analysis complete", 
                               f"Analyzed document and identified key components")
        
        # Step 2: Summary Generation
        self.update_agent_status("SummaryGenerator", "active", "Generating structured summary")
        
        summary_prompt = f"""
        You are an AI assistant designed to create client interaction summaries for Relationship Managers (RMs).
        
        Based on the document analysis provided below, generate a summary using this EXACT format (NO ** markdown formatting):

        Client Interaction Summary
        Date of Meeting: [Insert Date] 
        Participants: [List Participants] 
        Client Name: [Insert Client Name] 
        Meeting Type: [Call/Meeting/Other]

        1. Objectives of the Meeting :
        [Clearly state the purpose and goals of the meeting]

        2. Key Discussion Points : 
        [List the main topics, issues, and subjects discussed during the interaction]

        3. Decisions Made : 
        [Document any decisions, agreements, or resolutions reached during the meeting]

        4. Action Items :  
        [List specific action items with responsible for RM and Client separately and timelines where mentioned]

        Key Takeaways : 
        [Summarize the most important topics in bullet points, outcomes, or next steps from the meeting]

        Document Analysis Results:
        {analysis_result}

        Original Document Content:
        {document_text}

        Instructions:
        - Use the exact format structure shown above with the same headings but NO ** markdown formatting
        - Extract relevant information and place it under the appropriate sections
        - If information for any section is not available, note "[Information not available in document]"
        - Keep language professional and objective
        - Be concise while capturing all essential information
        - If any part is unclear or missing context, note this in the relevant section
        - Output clean, plain text without any ** formatting

        Generate the summary now following this exact format.
        """
        
        summary_messages = [{"role": "user", "content": summary_prompt}]
        summary_result = self.azure_client.generate_response(summary_messages)
        
        self.update_agent_status("SummaryGenerator", "complete", "Summary generation complete",
                               f"Generated structured summary with required format")
        
        # Step 3: Quality Review
        self.update_agent_status("QualityReviewer", "active", "Reviewing summary quality")
        
        review_prompt = f"""
        You are a Quality Reviewer for client interaction summaries. Review this summary to ensure it follows the exact required format and meets RM standards.
        
        Summary to review:
        {summary_result}
        
        Original document:
        {document_text}
        
        Required Format Check:
        The summary MUST follow this exact structure (WITHOUT ** markdown formatting):
        
        Client Interaction Summary
        Date of Meeting: [Insert Date] 
        Participants: [List Participants] 
        Client Name: [Insert Client Name] 
        Meeting Type: [Call/Meeting/Other]
        1. Objectives of the Meeting
        2. Key Discussion Points
        3. Decisions Made
        4. Action Items
        Key Takeaways
        
        Verify that the summary:
        1. Uses the exact format structure with proper headings but NO ** markdown formatting
        2. Has all required sections present
        3. Contains relevant information under each section
        4. Notes "[Information not available in document]" for missing information
        5. Maintains professional and objective language
        6. Is concise while capturing essential information
        7. Notes any unclear parts or missing context in relevant sections
        8. Does NOT contain any ** formatting - output should be clean plain text
        
        Provide the final, polished summary that strictly follows the required format WITHOUT any ** markdown formatting.
        """
        
        review_messages = [{"role": "user", "content": review_prompt}]
        final_summary = self.azure_client.generate_response(review_messages)
        
        self.update_agent_status("QualityReviewer", "complete", "Quality review complete",
                               f"Completed final review and provided polished summary")
        
        return {
            "analysis": analysis_result,
            "initial_summary": summary_result,
            "final_summary": final_summary,
            "processing_log": self.interaction_log.copy()
        }

def load_azure_config():
    """Load Azure OpenAI configuration from .env or Streamlit secrets"""
    config = {}
    
    # Try to load from .env file first
    if os.path.exists('.env'):
        from dotenv import load_dotenv
        load_dotenv()
        config = {
            'api_key': os.getenv('AZURE_OPENAI_API_KEY'),
            'endpoint': os.getenv('AZURE_OPENAI_ENDPOINT'),
            'api_version': os.getenv('AZURE_OPENAI_API_VERSION', '2024-02-15-preview'),
            'model': os.getenv('AZURE_OPENAI_MODEL', 'gpt-4')
        }
    else:
        # Fallback to Streamlit secrets
        try:
            config = {
                'api_key': st.secrets['AZURE_OPENAI_API_KEY'],
                'endpoint': st.secrets['AZURE_OPENAI_ENDPOINT'],
                'api_version': st.secrets.get('AZURE_OPENAI_API_VERSION', '2024-02-15-preview'),
                'model': st.secrets.get('AZURE_OPENAI_MODEL', 'gpt-4')
            }
        except KeyError:
            config = {}
    
    return config

def authenticate_user(username: str, password: str) -> bool:
    """Simple authentication function"""
    # Simple hash-based authentication (in production, use proper authentication)
    expected_hash = hashlib.sha256("aiplanet:aiplanet000".encode()).hexdigest()
    user_hash = hashlib.sha256(f"{username}:{password}".encode()).hexdigest()
    return user_hash == expected_hash

def show_login_screen():
    """Display login screen"""
    st.markdown("""
    <div class="login-container">
        <div class="login-header">🔐 Login to Client Interaction Summary Agent</div>
        <p>Please enter your credentials to access the Client Interaction Summary Generator</p>
    </div>
    """, unsafe_allow_html=True)
    
    with st.form("login_form"):
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            username = st.text_input("Username", placeholder="Enter username")
            password = st.text_input("Password", type="password", placeholder="Enter password")
            login_button = st.form_submit_button("🚀 Login", use_container_width=True)
            
            if login_button:
                if authenticate_user(username, password):
                    st.session_state.logged_in = True
                    st.session_state.username = username
                    st.rerun()
                else:
                    st.error("❌ Invalid credentials. Please try again.")
    
    # Add demo credentials info
    

def display_agent_status(agent_statuses: Dict[str, AgentStatus]):
    """Display current status of all agents"""
    st.subheader("🤖 Agent Status Dashboard")
    
    cols = st.columns(len(agent_statuses))
    
    for idx, (agent_name, status) in enumerate(agent_statuses.items()):
        with cols[idx]:
            status_color = {
                "active": "status-active",
                "waiting": "status-waiting", 
                "complete": "status-complete"
            }.get(status.status, "status-waiting")
            
            st.markdown(f"""
            <div class="agent-card">
                <h4><span class="status-indicator {status_color}"></span>{agent_name}</h4>
                <p><strong>Status:</strong> {status.status.title()}</p>
                <p><strong>Task:</strong> {status.current_task}</p>
            </div>
            """, unsafe_allow_html=True)

def display_agent_interactions(interaction_log: List[Dict]):
    """Display agent interactions in real-time"""
    if interaction_log:
        st.subheader("💬 Agent Interactions")
        
        for interaction in interaction_log[-5:]:  # Show last 5 interactions
            st.markdown(f"""
            <div class="agent-interaction">
                <p><strong>{interaction['timestamp']} - {interaction['agent']}:</strong></p>
                <p>{interaction['message']}</p>
                <small><em>Task: {interaction['task']}</em></small>
            </div>
            """, unsafe_allow_html=True)

def display_processing_step(step_name: str, description: str):
    """Display current processing step"""
    st.markdown(f"""
    <div class="processing-step">
        <h4>🔄 {step_name}</h4>
        <p>{description}</p>
    </div>
    """, unsafe_allow_html=True)

def save_to_history(filename: str, client_name: str, summary: str, full_results: Dict):
    """Save processing results to history"""
    if 'processing_history' not in st.session_state:
        st.session_state.processing_history = []
    
    history_item = ProcessingHistory(
        id=str(uuid.uuid4()),
        timestamp=datetime.now(),
        filename=filename,
        client_name=client_name,
        summary=summary,
        full_results=full_results
    )
    
    st.session_state.processing_history.insert(0, history_item)  # Add to beginning
    
    # Keep only last 50 items
    if len(st.session_state.processing_history) > 50:
        st.session_state.processing_history = st.session_state.processing_history[:50]

def display_history():
    """Display processing history in sidebar"""
    st.sidebar.header("📚 Processing History")
    
    if 'processing_history' not in st.session_state or not st.session_state.processing_history:
        st.sidebar.info("No processing history yet")
        return
    
    for i, item in enumerate(st.session_state.processing_history[:10]):  # Show last 10
        with st.sidebar.expander(f"📄 {item.filename[:20]}... ({item.timestamp.strftime('%m/%d %H:%M')})"):
            st.write(f"**Client:** {item.client_name}")
            st.write(f"**Date:** {item.timestamp.strftime('%Y-%m-%d %H:%M:%S')}")
            if st.button(f"Load Summary", key=f"load_{item.id}"):
                st.session_state.selected_history = item
                st.rerun()

def create_docx_summary(summary_text: str, filename: str) -> io.BytesIO:
    """Create a DOCX file from summary text"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Client Interaction Summary', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add content
    lines = summary_text.split('\n')
    for line in lines:
        if line.strip():
            if line.startswith('**') and line.endswith('**'):
                # Bold headers
                p = doc.add_paragraph()
                run = p.add_run(line.replace('**', ''))
                run.bold = True
            else:
                doc.add_paragraph(line)
    
    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def create_pdf_summary(summary_text: str, filename: str) -> io.BytesIO:
    """Create a PDF file from summary text"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        alignment=1,  # Center alignment
        spaceAfter=30
    )
    
    header_style = ParagraphStyle(
        'CustomHeader',
        parent=styles['Heading2'],
        fontSize=12,
        spaceBefore=12,
        spaceAfter=6
    )
    
    content = []
    content.append(Paragraph("Client Interaction Summary", title_style))
    content.append(Spacer(1, 12))
    
    lines = summary_text.split('\n')
    for line in lines:
        if line.strip():
            if line.startswith('**') and line.endswith('**'):
                # Bold headers
                content.append(Paragraph(line.replace('**', ''), header_style))
            else:
                content.append(Paragraph(line, styles['Normal']))
                content.append(Spacer(1, 6))
    
    doc.build(content)
    buffer.seek(0)
    return buffer

def extract_client_name_from_summary(summary: str) -> str:
    """Extract client name from summary"""
    lines = summary.split('\n')
    for line in lines:
        if 'Client Name:' in line:
            return line.split('Client Name:')[-1].strip().replace('[', '').replace(']', '')
    return "Unknown Client"

def main():
    """Main Streamlit application"""
    
    # Initialize session state
    if 'logged_in' not in st.session_state:
        st.session_state.logged_in = False
    if 'processing_history' not in st.session_state:
        st.session_state.processing_history = []
    
    # Show login screen if not logged in
    if not st.session_state.logged_in:
        show_login_screen()
        return
    
    # Logout button
    if st.button("🚪 Logout", key="logout", help="Logout"):
        st.session_state.logged_in = False
        st.session_state.clear()
        st.rerun()
    
    # Load Azure OpenAI configuration
    azure_config = load_azure_config()
    
    if not all(azure_config.values()):
        st.error("❌ Azure OpenAI configuration missing. Please set up your .env file or Streamlit secrets.")
        st.stop()
    
    # Title and header
    st.title("📋 Client Interaction Summary Generator")
    st.markdown(f"### Welcome back, {st.session_state.username}!")
    
    # Sidebar for configuration and history
    with st.sidebar:
        
        
        # Processing options
        st.header("📄 Processing Options")
        show_detailed_logs = st.checkbox("Show Detailed Agent Logs", value=True)
        auto_refresh = st.checkbox("Auto-refresh Agent Status", value=True)
        
        st.markdown("---")
        
        # Display history
        display_history()
    
    # Check if a history item is selected
    if 'selected_history' in st.session_state:
        st.success(f"📄 Loaded from history: {st.session_state.selected_history.filename}")
        
        # Display the historical summary
        st.markdown(f"""
        <div class="summary-output">
            <h3>📊 Historical Client Interaction Summary</h3>
            {st.session_state.selected_history.summary.replace(chr(10), '<br>')}
        </div>
        """, unsafe_allow_html=True)
        
        # Download options for historical summary
        st.subheader("💾 Download Historical Summary")
        col1, col2 = st.columns(2)
        
        with col1:
            docx_buffer = create_docx_summary(
                st.session_state.selected_history.summary, 
                st.session_state.selected_history.filename
            )
            st.download_button(
                "📄 Download as DOCX",
                data=docx_buffer.getvalue(),
                file_name=f"summary_{st.session_state.selected_history.timestamp.strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        
        with col2:
            pdf_buffer = create_pdf_summary(
                st.session_state.selected_history.summary, 
                st.session_state.selected_history.filename
            )
            st.download_button(
                "📄 Download as PDF",
                data=pdf_buffer.getvalue(),
                file_name=f"summary_{st.session_state.selected_history.timestamp.strftime('%Y%m%d_%H%M%S')}.pdf",
                mime="application/pdf"
            )
        
        if st.button("🔄 Process New Document"):
            del st.session_state.selected_history
            st.rerun()
        
        return
    
    # Main content area
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.header("📤 Document Upload")
        
        # File upload
        uploaded_file = st.file_uploader(
            "Upload Client Interaction Document",
            type=['pdf', 'docx', 'txt'],
            help="Supported formats: PDF, DOCX, TXT"
        )
        
        # Format template (optional)
        st.subheader("📝 Summary Format Template (Optional)")
        format_template = st.text_area(
            "Enter custom format template:",
            placeholder="e.g., Meeting Overview, Key Discussions, Action Items, Next Steps",
            height=100
        )
    
    with col2:
        st.header("🎯 Processing Controls")
        
        # Processing button
        process_button = st.button(
            "🚀 Generate Summary", 
            type="primary",
            disabled=not uploaded_file
        )
        
        if not uploaded_file:
            st.info("📁 Please upload a document to process")
    
    # Processing section
    if process_button and uploaded_file:
        
        # Initialize Azure OpenAI client
        try:
            azure_client = AzureOpenAIWrapper(
                api_key=azure_config['api_key'],
                endpoint=azure_config['endpoint'],
                api_version=azure_config['api_version'],
                model=model_choice
            )
        except Exception as e:
            st.error(f"Failed to initialize Azure OpenAI client: {str(e)}")
            st.stop()
        
        # Initialize the multi-agent system
        summary_agents = ClientSummaryAgents(azure_client)
        
        # Document processing
        display_processing_step("Document Reading", "Extracting text from uploaded document...")
        
        # Extract text based on file type
        if uploaded_file.type == "application/pdf":
            document_text = DocumentProcessor.extract_text_from_pdf(uploaded_file)
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            document_text = DocumentProcessor.extract_text_from_docx(uploaded_file)
        else:
            document_text = DocumentProcessor.extract_text_from_txt(uploaded_file)
        
        if document_text.strip():
            st.success(f"✅ Successfully extracted {len(document_text)} characters from document")
            
            # Show document preview
            with st.expander("📖 Document Preview"):
                st.text(document_text[:500] + "..." if len(document_text) > 500 else document_text)
            
            # Agent processing
            display_processing_step("Agent Processing", "Multi-agent system is analyzing and summarizing...")
            
            # Create placeholders for dynamic updates
            status_placeholder = st.empty()
            interaction_placeholder = st.empty()
            
            # Process document through agents
            with st.spinner("🤖 Agents are working on your document..."):
                result = summary_agents.process_document(document_text, format_template)
            
            # Display final results
            st.success("🎉 Summary generation completed!")
            
            # Final agent status
            with status_placeholder.container():
                display_agent_status(summary_agents.agent_statuses)
            
            # Agent interactions
            if show_detailed_logs:
                with interaction_placeholder.container():
                    display_agent_interactions(summary_agents.interaction_log)
            
            # Extract client name for history
            client_name = extract_client_name_from_summary(result['final_summary'])
            
            # Save to history
            save_to_history(
                filename=uploaded_file.name,
                client_name=client_name,
                summary=result['final_summary'],
                full_results=result
            )
            
            # Display final summary
            st.markdown(f"""
            <div class="summary-output">
                <h3>📊 Final Client Interaction Summary</h3>
                {result['final_summary'].replace(chr(10), '<br>')}
            </div>
            """, unsafe_allow_html=True)
            
            # Download options
            st.subheader("💾 Download Options")
            
            col1, col2 = st.columns(2)
            
            with col1:
                # Create DOCX
                docx_buffer = create_docx_summary(result['final_summary'], uploaded_file.name)
                st.download_button(
                    "📄 Download as DOCX",
                    data=docx_buffer.getvalue(),
                    file_name=f"client_summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            
            with col2:
                # Create PDF
                pdf_buffer = create_pdf_summary(result['final_summary'], uploaded_file.name)
                st.download_button(
                    "📄 Download as PDF",
                    data=pdf_buffer.getvalue(),
                    file_name=f"client_summary_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                    mime="application/pdf"
                )
        
        else:
            st.error("❌ Failed to extract text from the document. Please check the file format and try again.")
    
    # Footer
    st.markdown("---")
    st.markdown("""
    <div style="text-align: center; color: #666;">
        <p>🔒 Your documents are processed securely and not stored permanently.</p>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
